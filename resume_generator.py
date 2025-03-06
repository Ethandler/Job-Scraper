import json
import os
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document

# 🔹 Load User Data
def load_user_data():
    """Load user details from JSON files."""
    with open("user_data.json", "r") as file:
        user_data = json.load(file)

    with open("user_profile.json", "r") as file:
        user_profile = json.load(file)

    return user_data, user_profile

# 🔹 Backup Existing Resumes
def backup_old_resumes():
    """Back up old resume files before generating new ones."""
    for ext in ["pdf", "docx"]:
        old_file = f"resume.{ext}"
        backup_file = f"resume_backup.{ext}"
        
        if os.path.exists(old_file):
            if os.path.exists(backup_file):  # Remove old backup if it exists
                os.remove(backup_file)
            os.rename(old_file, backup_file)  # Rename old resume

# 🔹 Generate PDF Resume
def create_pdf_resume(user_data, user_profile, output_file="resume.pdf"):
    """Generate a PDF resume from user data."""
    c = canvas.Canvas(output_file, pagesize=letter)
    width, height = letter

    # 🔹 Title
    c.setFont("Helvetica-Bold", 18)
    c.drawString(50, height - 50, f"{user_data['first_name']} {user_data['last_name']}")

    # 🔹 Contact Details
    c.setFont("Helvetica", 12)
    c.drawString(50, height - 80, f"📧 {user_data['email']}")
    c.drawString(50, height - 100, f"📞 {user_data['phone']}")
    c.drawString(50, height - 120, f"📍 {user_data['location']}")
    
    if user_data.get("linkedin") and user_data["linkedin"] != "Not provided":
        c.drawString(50, height - 140, f"🔗 LinkedIn: {user_data['linkedin']}")
    if user_data.get("github") and user_data["github"] != "Not provided":
        c.drawString(50, height - 160, f"💻 GitHub: {user_data['github']}")

    # 🔹 Career Summary
    y_position = height - 200
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y_position, "💼 Career Summary")
    y_position -= 20
    c.setFont("Helvetica", 12)
    c.drawString(50, y_position, f"Preferred Roles: {user_profile['preferred_job_titles']}")

    # 🔹 Skills
    y_position -= 30
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y_position, "🚀 Skills")
    y_position -= 20
    c.setFont("Helvetica", 12)

    for skill in user_profile["skills"].split(", "):
        c.drawString(50, y_position, f"• {skill}")
        y_position -= 15

    # 🔹 Work Experience
    y_position -= 30
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y_position, "📜 Work Experience")
    y_position -= 20
    c.setFont("Helvetica", 12)

    for job in user_profile["work_experience"]:
        c.drawString(50, y_position, f"{job['title']} at {job['company']} ({job['years']} years)")
        y_position -= 15
        c.drawString(50, y_position, job["description"])
        y_position -= 25

    # 🔹 Save PDF
    c.save()
    print(f"✅ Resume PDF generated: {output_file}")

# 🔹 Generate Word Resume (DOCX)
def create_docx_resume(user_data, user_profile, output_file="resume.docx"):
    """Generate a Word (DOCX) resume from user data."""
    doc = Document()

    # 🔹 Title
    doc.add_heading(f"{user_data['first_name']} {user_data['last_name']}", level=1)

    # 🔹 Contact Details
    doc.add_paragraph(f"📧 {user_data['email']} | 📞 {user_data['phone']} | 📍 {user_data['location']}")
    if user_data.get("linkedin") and user_data["linkedin"] != "Not provided":
        doc.add_paragraph(f"🔗 LinkedIn: {user_data['linkedin']}")
    if user_data.get("github") and user_data["github"] != "Not provided":
        doc.add_paragraph(f"💻 GitHub: {user_data['github']}")

    # 🔹 Career Summary
    doc.add_heading("💼 Career Summary", level=2)
    doc.add_paragraph(f"Preferred Roles: {user_profile['preferred_job_titles']}")

    # 🔹 Skills
    doc.add_heading("🚀 Skills", level=2)
    for skill in user_profile["skills"].split(", "):
        doc.add_paragraph(f"• {skill}")

    # 🔹 Work Experience
    doc.add_heading("📜 Work Experience", level=2)
    for job in user_profile["work_experience"]:
        doc.add_paragraph(f"{job['title']} at {job['company']} ({job['years']} years)")
        doc.add_paragraph(job["description"])

    # 🔹 Save DOCX
    doc.save(output_file)
    print(f"✅ Resume DOCX generated: {output_file}")

# 🔹 Run Resume Generator
if __name__ == "__main__":
    print("\n🔹 Backing up old resumes...")
    backup_old_resumes()
    
    print("\n🔹 Generating new resume...")
    user_data, user_profile = load_user_data()
    create_pdf_resume(user_data, user_profile)
    create_docx_resume(user_data, user_profile)

    print("\n✅ Resume generation complete!")
